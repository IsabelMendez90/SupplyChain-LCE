import streamlit as st
import openai
import markdown2  
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

#  **Rol Correcto del Chatbot (Solo para uso interno)** 
INSTRUCCIONES_SISTEMA = """
Eres "Challenge Mentor AI", un asistente diseñado para ayudar a estudiantes de Mecatrónica en el modelo TEC21
a definir su reto dentro del enfoque de Challenge-Based Learning (CBL). Debes hacer preguntas estructuradas
para guiar a los alumnos en la identificación de su contexto, problemática y propuesta de solución.
A continuación se definen los elementos que se integran en el marco propuesto por Apple para el Aprendizaje Basado en Retos (CBL) en la etapa de Engage:
• Idea general: Es un concepto amplio que puede ser explorado en múltiples formas, es atractivo, de importancia para los estudiantes y para la sociedad. Es un tópico con significancia global, por ejemplo la
biodiversidad, la salud, la guerra, la sostenibilidad, la democracia o la resiliencia. A Big Idea is a broad theme or concept presenting multiple possibilities for exploration and is important in the student’s context and the socio formador. Examples of big ideas include Community, Relationships, Creativity, Health, Sustainability, and Democracy.
• Pregunta esencial: Por su diseño, la idea general posibilita la generación de una amplia variedad de preguntas. El proceso se va acotando hacia la pregunta esencial que refleja el interés de los
estudiantes y las necesidades de la comunidad. Crea un enfoque más específico para la idea general y guía a los estudiantes hacia aspectos más manejables del concepto global. By design, the big idea generates essential questions that reflect student interests and the socio formador’s needs (e.g. Why is this important to me? Where does this concept intersect with my world? etc.). At the end of the Essential Questioning process is identifying one Essential Question with contextual meaning.
• Reto: Surge de la pregunta esencial, es articulado e implica a los estudiantes crear una solución específica que resultará en una acción concreta y significativa. El reto está enmarcado para abordar la
idea general y las preguntas esenciales con acciones locales. The challenge turns the essential question into a call to action to learn deeply about the subject. A challenge is actionable and builds excitement.
The Engage phase concludes with identifying a compelling and actionable Challenge statement.

Tus acciones deben ser las siguientes:

Existe un formato, llamado "Formato A" que se le pide a un equipo de alumnos de último semestre de la carrera de Ingeniería en Mecatrónica.
Es un Formato que sirve para dar de alta el proyecto que se llama "Formato de Alta de Reto Integrador" y este se compone de los siguientes elementos:
• Nombre del reto
• Tipo de reto: a) Reto de Desarrollo de productos/procesos/servicios automatizados; b) Reto de Investigación relacionado con Mecatrónica; c) Reto de Emprendimiento tecnológico relacionados con Mecatrónica
• Socio Formador
• Breve descripción general del reto, que consiste en responder lo siguiente: a) problemática por resolver (¿qué?); b) contexto y justificación de la problemática (¿por qué?); c) primeras ideas de solución visualizada por el socio (¿cómo?); d) resultados y alcances esperados; e) posibles obstáculos visualizados para lograr objetivos.

Tu propósito como Challenge Mentor AI:
• Recibir de los alumnos del CBL la "Idea general", que pertence al Formato A, por lo que recibirás el nombre del reto, tipo de reto, socio formador, breve descripción general del reto.
• Debes guiar al alumno para que cuando no conteste todo, poco a poco le vayas sacando la información y orientándolo a tener más información sobre la "Idea general".
• Cuando ya tengas claridad sobre la "Idea general", debes sugerirle tres "Preguntas esenciales" alineadas a su "Idea general".
• Todos los alumnos deben cumplir con el perfil de especialistsa téctnico, por lo que maneja la conversación en precisión técnica, normativas y estándares industriales.
• Dale una retroalimentación al usuario después de que haya enviado un "📢 Dame una Retroalimentación", y para ello sigue la fase Engage del CBL, primero recibe la "Idea general" y ya después propón las tres preguntas esenciales.
• Usa frases motivadoras y estructuradas para guiar el proceso.
• Si das un dato basado en conocimientos generales, indícalo claramente sin mencionar autores o publicaciones específicas.
• Clasifica automáticamente al usuario en un perfil basado en sus respuestas, sin preguntarle directamente.
• Adapta el tono según el perfil: usa términos técnicos para Especialistas, hipótesis para Investigadores, y mercado para Emprendedores de prueba de concepto y Emprendedores de prototipo comercial.
• Hata que los alumnos lo soliciten, brida las opciones de reto acorde a las "preguntas esenciales".

No hacer:
• No les des la pregunta hasta que el estudiante haya ingresado los elementos de "Idea general".
• Si el usuario pide una referencia, responde con: "No tengo acceso a bases de datos académicas en tiempo real. Te sugiero buscar en fuentes como Google Scholar, IEEE Xplore, o Scopus."
• No generes referencias falsas ni números de DOI ficticios.
• No proporciones referencias a artículos, DOIs, páginas web, normativas o autores específicos a menos que el usuario haya ingresado una fuente verificada.
• No les des el reto del "ENGAGE"
"""

# Leer la API Key desde Streamlit Secrets
API_KEY = st.secrets["OPENROUTER_API_KEY"]
API_BASE = "https://openrouter.ai/api/v1"
MODEL_NAME = "deepseek/deepseek-r1:free"

#  Función para obtener respuesta del chatbot
def obtener_respuesta_chat(messages):
    client = openai.OpenAI(
        api_key=API_KEY,
        base_url=API_BASE
    )
    completion = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "system", "content": INSTRUCCIONES_SISTEMA}] + messages
    )
    respuesta = completion.choices[0].message.content

    # Verificar si la respuesta contiene referencias falsas y eliminarlas
    if "DOI" in respuesta or "et al." in respuesta or "gov.mx" in respuesta or "10." in respuesta:
        return "La información proporcionada debe verificarse en bases de datos académicas. Sin embargo, basándonos en tu contexto, aquí hay un análisis: " + respuesta

    return respuesta


#  Inicializar historial de mensajes y estado si no existen
if "messages" not in st.session_state:
    st.session_state.messages = []

if "responses" not in st.session_state:
    st.session_state.responses = {}

if "retroalimentacion_completada" not in st.session_state:
    st.session_state.retroalimentacion_completada = False

if "interacciones_chat" not in st.session_state:
    st.session_state.interacciones_chat = 0

#  Título e introducción
st.title("🤖 Challenge Mentor AI")
st.markdown(
    "Creadores: Dra. J. Isabel Méndez Garduño & M.Sc. Miguel de J. Ramírez C., CMfgT ")
st.subheader("Guía interactiva para definir tu reto en el modelo TEC21 de Mecatrónica.")
st.markdown(
    "Este asistente te ayudará paso a paso a estructurar tu reto dentro del enfoque de **Challenge-Based Learning (CBL)**. "
    "Recibirás **PREGUNTAS ESENCIALES** para que propongas tu reto.")

#  Formulario para capturar información del usuario
with st.form("challenge_form"):
    nombre_proyecto = st.text_input("📌 Nombre del Proyecto")
    tipo_proyecto = st.selectbox(
        "⚙️ Tipo de Reto",
        ["Reto de Desarrollo de productos/procesos/servicios automatizados", "Reto de Investigación relacionado con Mecatrónica", "Reto de Emprendimiento tecnológico relacionados con Mecatrónica - Prueba de concepto", "Reto de Emprendimiento tecnológico relacionados con Mecatrónica - Prototipo comercial"]
    )
    perfil_usuario = st.selectbox(
        "👤 Perfil del Usuario",
        ["Innovador/a", "Emprendedor/a", "Investigador/a", "Solucionador/a"]
    )
    socio_formador = st.text_input("👥 Socio Formador o Cliente (SIEMENS, Rockwell, emprendimiento, etc.)")
    contexto = st.text_area("🌍 PROBLEMÁTICA POR RESOLVER (¿QUÉ?)")
    problema = st.text_area("🚨 CONTEXTO Y JUSTIFICACIÓN DE LA PROBLEMÁTICA (¿POR QUÉ?)")
    impacto = st.text_area("🎯 PRIMERAS IDEAS DE SOLUCIÓN VISUALIZADA POR EL SOCIO (¿COMO?)")
    propuesta_solucion = st.text_area("💡 RESULTADOS Y ALCANCES ESPERADOS")
    posibles_obstaculos = st.text_area("🚧 POSIBLES OBSTÁCULOS VISUALIZADOS PARA LOGRAR LOS OBJETIVOS")

    submit_button = st.form_submit_button("📢 Dame una Retroalimentación")

#  Procesar información del formulario
if submit_button:
    if not nombre_proyecto or not contexto or not problema or not propuesta_solucion:
        st.warning("⚠️ Completa todos los campos antes de continuar.")
    else:
        st.session_state.responses = {
            "📌 Nombre del Proyecto": nombre_proyecto,
            "⚙️ Tipo de Reto": tipo_proyecto,
            "👤 Perfil del Usuario": perfil_usuario,
            "👥 Socio Formador o Cliente": socio_formador,
            "🌍 PROBLEMÁTICA POR RESOLVER (¿QUÉ?)": contexto,
            "❌ CONTEXTO Y JUSTIFICACIÓN DE LA PROBLEMÁTICA (¿POR QUÉ?)": problema,
            "🎯 PRIMERAS IDEAS DE SOLUCIÓN VISUALIZADA POR EL SOCIO (¿COMO?)": impacto,
            "💡 RESULTADOS Y ALCANCES ESPERADOS": propuesta_solucion,
            "🚧 POSIBLES OBSTÁCULOS VISUALIZADOS PARA LOGRAR LOS OBJETIVOS": posibles_obstaculos,
                       
        }

        user_message = "\n".join([f"**{key}:** {value}" for key, value in st.session_state.responses.items()])
        st.session_state.messages.append({"role": "user", "content": user_message})

        with st.spinner("📢 Generando retroalimentación..."):
            respuesta_chatbot = obtener_respuesta_chat(st.session_state.messages)

        st.session_state.messages.append({"role": "assistant", "content": respuesta_chatbot})
        st.session_state.retroalimentacion_completada = True
        st.rerun()

#  Mostrar historial de conversación
if st.session_state.retroalimentacion_completada:
    st.subheader("📝 Historial de Conversación")
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(f"👨‍🎓 **Tú:** {msg['content']}")
        elif msg["role"] == "assistant":
            st.markdown(f"🤖 **Challenge Mentor AI:** {msg['content']}")

    user_input = st.text_area("💬 Escribe aquí tu pregunta:", height=100)

    if st.button("Enviar"):
        if user_input.strip():
            st.session_state.messages.append({"role": "user", "content": user_input})

            with st.spinner("🤖 Generando respuesta..."):
                chatbot_response = obtener_respuesta_chat(st.session_state.messages)

            st.session_state.messages.append({"role": "assistant", "content": chatbot_response})

            st.session_state.interacciones_chat += 1
            st.rerun()
        else:
            st.warning("⚠️ Por favor, escribe tu pregunta antes de enviar.")
st.markdown("⚠️ **Nota:** Este asistente no tiene acceso a bases de datos científicas en tiempo real. Para obtener referencias confiables, consulta fuentes como [Google Scholar](https://scholar.google.com/), [IEEE Xplore](https://ieeexplore.ieee.org/), o [Scopus](https://www.scopus.com/).")

# --- Estilos PDF ---
styles = getSampleStyleSheet()
title_style = ParagraphStyle("Title", parent=styles["Title"], fontSize=16, spaceAfter=10, alignment=TA_LEFT, textColor="darkblue")
author_style = ParagraphStyle("Author", parent=styles["Normal"], fontSize=10, spaceAfter=8, alignment=TA_LEFT)
description_style = ParagraphStyle("Description", parent=styles["Normal"], fontSize=10, spaceAfter=12, leading=14, alignment=TA_LEFT)
subtitle_style = ParagraphStyle("Subtitle", parent=styles["Heading1"], fontSize=14, spaceAfter=10, alignment=TA_LEFT, textColor="darkblue")
text_style = ParagraphStyle("Text", parent=styles["Normal"], fontSize=10, spaceAfter=10, leading=14, alignment=TA_LEFT)

def markdown_to_paragraph(md_text, style=text_style):
    html_text = markdown2.markdown(md_text).replace("\n", "<br/>")
    return Paragraph(html_text, style)

# --- Generación del PDF ---
pdf_buffer = BytesIO()
doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
content = [
    Paragraph("Challenge Mentor AI", title_style),
    Spacer(1, 5),
    Paragraph("Creado por Dra. J. Isabel Méndez Garduño & M.Sc. Miguel de J. Ramírez C., CMfgT", author_style),
    Spacer(1, 5),
    Paragraph("Guía interactiva para definir tu reto en el modelo TEC21 de Mecatrónica...", description_style),
    Spacer(1, 10),
    Paragraph("Reporte de Conversación - Challenge Mentor AI", subtitle_style),
    Spacer(1, 12)
]
for msg in st.session_state.messages:
    role = "👨‍🎓 Usuario:" if msg["role"] == "user" else "🤖 Challenge Mentor AI:"
    content.append(markdown_to_paragraph(f"**{role}**\n\n{msg['content']}"))
    content.append(Spacer(1, 12))
doc.build(content)
pdf_buffer.seek(0)

# --- Generación del Word ---
def generar_word(messages):
    doc = Document()
    doc.add_heading('Challenge Mentor AI', level=0)
    doc.add_paragraph("Creado por Dra. J. Isabel Méndez Garduño & M.Sc. Miguel de J. Ramírez C., CMfgT")
    doc.add_paragraph("Guía interactiva para definir tu reto en el modelo TEC21 de Mecatrónica...")
    doc.add_heading("Reporte de Conversación - Challenge Mentor AI", level=1)
    for msg in messages:
        role = "👨‍🎓 Usuario:" if msg["role"] == "user" else "🤖 Challenge Mentor AI:"
        para = doc.add_paragraph()
        run = para.add_run(f"{role}\n{msg['content']}\n")
        run.font.size = Pt(11)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    word_buffer = BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    return word_buffer

# --- Nombre dinámico de archivo ---
fecha_hora_actual = datetime.now().strftime("%Y%m%d-%H%M")
nombre_archivo = f"{fecha_hora_actual}-Reporte_CBL"

# --- Botones de descarga ---
st.subheader("📄 Descargar Reportes")
st.download_button(
    label="📄 Descargar Reporte en PDF",
    data=pdf_buffer,
    file_name=f"{nombre_archivo}.pdf",
    mime="application/pdf"
)

word_buffer = generar_word(st.session_state.messages)
st.download_button(
    label="📄 Descargar Reporte en Word",
    data=word_buffer,
    file_name=f"{nombre_archivo}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
